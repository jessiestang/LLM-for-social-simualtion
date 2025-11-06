from openai import OpenAI
from docx import Document
import json
import os
import re


class ModelConstructor:
    def __init__(self, model_name="gpt-4o-mini"):
        self.client = OpenAI(
            api_key=os.getenv("OPENAI_API_KEY")
        )  # set your API key in environment variable
        self.model_name = model_name  # default model
    
    def model_brainstorming(self, problem_context: str):
        """
        This llm agent will do three tasks:
        (1) reconstruct the classic agent-based model based on the problem context provided by users;
        (2) propose alternatives or novel variants of the classic model;
        (3) validate suggestions by literature (to be implemented in future).
        """
        with open(problem_context, "r", encoding="utf-8") as file:
            problem_definition = file.read()
        
        # prompt the LLM to brainstorm the actions of agents, and the motivations behind the actions
        system_prompt = """
        You are a researcher in computational social science and agent-based modeling.
        You secialize in brainstorming how classic agent-based models can be extended to capture new social mechanisms.
        You will be provided with a problem context.
        You have three tasks. Please accomplish them step-by-step:
        (1) Reconstruct the traditional agent-based model that best fits the problem context.
        (2) Propose at least three novel extensions of the traditional model to better capture the social mechanisms in the problem context.
        (3) For each proposed extension, provide literature references that support your suggestions.

        For each version of model, please describe:
        (1) key agents and their attributes;
        (2) environment and interaction structure;
        (3) expected emergent behaviors;
        (4) theoretical justifications of extended rules from literature.
        (5) why this new variant can yield new insights

        Output strictly in JSON format:
        [
          {
            "model_version": "traditional" or "extension_1" or "extension_2" or "extension_3",
            "key_agents": ["agent1", "agent2", "..."],
            "agent_attributes": ["attribute1", "attribute2", "..."],
            "environment": "description of the environment",
            "interaction_structure": "description of interaction structure",
            "expected_emergent_behaviors": ["behavior1", "behavior2", "..."],
            "theoretical_justifications": ["justification1", "justification2", "..."],
            "literature_references": ["reference1", "reference2", "..."],
            "insightfulness": "description of why this variant can yield new insights" #  extension only
          },
            ...
        ]
        """
        user_prompt = f"""The given problem context is:
        {problem_definition}
        Please brainstorm the model construction based on the three tasks mentioned above.
        """
        # call the LLM model
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        # parse the response to extract JSON
        pb = response.choices[0].message.content
        try:
            rules = json.loads(pb)
            return rules
        except json.JSONDecodeError:
            print(
                "The response is not valid JSON. Please check the output format.")
            return pb
    
    def extract_problem_definition(self, file_path: str):
        """
        This llm agent reads the problem definition provided by users and extracts key information in json format
        """
        # load problem definition given by users
        with open(file_path, "r", encoding="utf-8") as file:
            problem_definition = file.read()

        # prompt the LLM to extract key information
        system_prompt = """
        You are a research assistant who specializes in social simulation models.
        You are given a problem definition provided by users.
        Your task is to extract key information from the problem definition, including:
        1. types and roles of involved agents
        2. environment and interaction structure
        3. key behaviors and decision-making processes of agents
        Output strictly in JSON format:
        {
          "agent_types": ["type1", "type2", "..."],
          "environment": "description of the environment",
          "interaction_structure": "description of interaction structure",
          "key_behaviors": ["behavior1", "behavior2", "..."],
          "decision_making_processes": ["process1", "process2", "..."],
          ""key parameters": ["parameter1", "parameter2", "..."]
        }
        Ensure your response is valid JSON that can be parsed by a JSON parser.

        Example
        The given problem definition is:
        "We want to model the voting behaviour of citizens in a democratic society.
        Suppose there are two types of voters with different political preferences (+1 and -1).
        Voters interact in a social network where they can influence each other's opinions.
        Each voter decides whether to keep or change their opinin based on that of their neighbours."

        The extracted key information should be:
        {
          "agent_types": ["type1: preference +1", "type2: preference -1"],
          "environment": "Voters are placed in a social network",
          "interaction_structure (among agents)": "voters can read opinions from their neighbours and influence each other",
          "interaction_structure (with environment)": "N/A",
          "key_behaviors": ["interacting with neighbours", "updating own opinion"],
          "decision_making_processes": ["deciding whether to keep or change opinion based on neighbours' opinions"],
          "key parameters": ["network structure", "initial opinion distribution", "threshold for opinion change"]
        }
        """
        user_prompt = problem_definition

        # call the LLM model
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        # parse the response to extract JSON
        pb = response.choices[0].message.content
        try:
            features = json.loads(pb)
        except json.JSONDecodeError:
            raise ValueError(
                "The response is not valid JSON. Please check the output format."
            )
        return features
    
    def model_generation_refining(self, features: str):
        """
        This llm agents refine the brainstorming ideas and transform it into concrete description of the model
        """
        # prompt the llm to generate agent rules
        system_prompt = """
        You are a helpful research assistant who specializes in constructing social simulation models based on provided context.
        You will be provided with a short context of an agent-based model.
        Your task is to dive deeper into the context and enlarge the context into a specific conceptual model for agent-based modeling.
        Please strictly stick to the provided context and do not add any extra assumptions!!

        You should follow these steps:
        (1) Think about how the agents will be initialized. What attributes will they have? What actions can they perform?
        (2) Consider the environment in which the agents operate. How is it structured? How do agents interact with each other and with the environment?
        (3) Define the decision-making processes of the agents. How do they decide what actions to take based on their attributes and the state of the environment?
        (4) Specify the rules that govern agent behavior. What conditions lead to specific actions? How do agents adapt or learn over time?

        CRITICAL: Output ONLY valid JSON. Do not include any text before or after the JSON. Do not wrap in markdown code blocks.
        Start your response directly with { and end with }.

        Your output should be in JSON format, clearly outlining the agent rules and model structure:
        {
            "agent_initialization": {
                "attributes": ["attribute1", "attribute2", "..."],
                "actions": ["action1", "action2", "..."],
                "agent_types": ["type1", "type2", "..."],
                "initial_distribution": "description of how agents are initially distributed",
                "size of agents": "number of agents in the model",
                "adaptation_mechanisms": "description of how agents adapt or learn over time"
            },
            "environment": {
                "structure": "description of the environment structure",
                "interactions": "description of agent-agent and agent-environment interactions",
                "changes_over_time": "description of how the environment changes over time, and how agents influence these changes",
                "key parameters": ["parameter1", "parameter2", "..."],
            },
            "decision_making_processes": ["process1", "process2", "..."],
            "behavior_rules": [
                "Rule 1: specific behavior description",
                "Rule 2: specific behavior description",
                "Rule 3: "...",
                "..."
            ],
            "expected_emergent_behaviors": [
                "Behavior 1: what emerges",
                "Behavior 2: what emerges",
                "..."
            ]
            }
        Ensure your response is valid JSON that can be parsed by a JSON parser.
        """

        user_prompt = f"""
        please construct a very detailed agent-based model based on this brainstorming idea {json.dumps(features, indent = 2)}.
        Follow the steps mentioned in the system prompt to generate a comprehensive model description with clear agent rules.
        Output ONLY the JSON object, no other text.
        Please strictly stick to the provided context and do not add any extra assumptions!!
        """

        # call the LLM model
        llm_model = self.client.chat.completions.create(
            model=self.model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )

        # parse the json response
        response = llm_model.choices[0].message.content
        try:
            rules = json.loads(response)
            if isinstance(rules, dict):
                rules = [rules]  # Ensure it's a list of models
            elif isinstance(rules, list):
                pass
            else:
                raise ValueError("The response is not a valid list or dictionary.")
    
        except json.JSONDecodeError:
            match = re.search(r'\[.*\]', response, re.DOTALL)
            if match:
                rules = json.loads(match.group())
                if isinstance(rules, dict):
                    rules = [rules]  # Ensure it's a list of models
                elif isinstance(rules, list):
                    pass
            else:
                raise ValueError("LLM output not valid JSON, here’s raw output:")
        return rules

    def refine_with_feedback(self, agent_rules: str, user_feedback: str):
        """
        This LLM agent reflects on the generated model rules and suggests improvements. It also takes users's feedback into account
        """
        # load user input
        system_prompt = """You are a research assistant who specializes in social simulation models. You are given a description of agent-based model and user feedback.
        Your task is to refine the conceptual model based on the feedback and your own reflection. Output the refined model rules in JSON format."""
        user_prompt = f"""The given conceptual model is: {agent_rules}. The user feedback is: {user_feedback}. 
        Please refine the model rules based on the feedback and your own reflection. Output the refined model rules in JSON format."""

        # call the LLM model
        llm_model = self.client.chat.completions.create(
            model=self.model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )

        # parse the json response
        response = llm_model.choices[0].message.content
        try:
            refined_rules = json.loads(response)
            return json.dumps(refined_rules, indent=2)
        except json.JSONDecodeError:
            print("LLM output not valid JSON, here’s raw output:")
            return response

    def model_construction_pipeline(self, file_path: str, save_path: str):
        """
        This function provides a pipeline that:
        (1) brainstorming model ideas based on problem definition;
        (2) generate conceptual model based on brainstorming ideas;
        (3) refine the model based on user feedback;
        (4) save and export the conceptual model.
        """
        # first get the features
        print("Step 1: Brainstorming model ideas based on problem definition...")
        model_ideas = self.model_brainstorming(file_path)
        print("Brainstormed model ideas:")
        print(json.dumps(model_ideas, indent=2))
        idea_type = int(input("which model do you want to choose (1/2/3)?"))
        selected_idea = model_ideas[idea_type - 1]
        #TODO: enable human feedback here to regenerate ideas if needed

        # generate the agent rules
        print("Step 2: Generating the conceptual model based on the selected idea")
        conceptual_model = self.model_generation_refining(selected_idea)
        print("Generated conceptual model:")
        print(json.dumps(conceptual_model, indent=2, ensure_ascii=False))

        user_input = input("do you want to provide any feedback (y/n)?")
        refined_model = None
        while user_input.lower() == "y":
            user_feedback = input("please provide your feedback:")
            print("Step 3: Refining model based on user feedback...")
            refined_model = self.refine_with_feedback(conceptual_model, user_feedback)
            print("Refined conceptual model:")
            print(refined_model)
            user_input = input("do you want to provide any feedback (y/n)?")

        # save and export the conceptual model
        print("Step 4: Saving and exporting the conceptual model...")
        if refined_model:
            with open(save_path, "w", encoding="utf-8") as f:
                json.dump(refined_model, f, indent=2, ensure_ascii=False)  # save json output for code assistant
            #self.save_to_wordfile(refined_model, save_path.replace(".json", ".docx"))
        else:
            with open(save_path, "w", encoding="utf-8") as f:
                json.dump(conceptual_model, f, indent=2, ensure_ascii=False)  # save json output for code assistant
            #self.save_to_wordfile(conceptual_model, save_path.replace(".json", ".docx"))  # save word file for users

    def save_to_wordfile(self, model_description: str, file_path: str):
        """
        This function saves the model description and export it as a word file
        """
        doc = Document()
        doc.add_heading("Conceptual Model Description", level=1)

        # Handle both string and dict/list inputs
        if isinstance(model_description, str):
            sections = json.loads(model_description)
        elif isinstance(model_description, (dict, list)):
            sections = model_description
        else:
            raise ValueError(f"Unexpected type for model_description: {type(model_description)}")
        
            # Handle list of models
        if isinstance(sections, list):
            for i, model in enumerate(sections, 1):
                doc.add_heading(f"Model {i}", level=1)
                for section, content in model.items():
                    doc.add_heading(section.replace("_", " ").title(), level=2)
                    
                    # Handle nested structures
                    if isinstance(content, dict):
                        for key, value in content.items():
                            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
                    elif isinstance(content, list):
                        for item in content:
                            doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        doc.add_paragraph(str(content))
                
                if i < len(sections):
                    doc.add_page_break()
        
        # Handle single model (dict)
        else:
            for section, content in sections.items():
                doc.add_heading(section.replace("_", " ").title(), level=2)
                
                # Handle nested structures
                if isinstance(content, dict):
                    for key, value in content.items():
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
                elif isinstance(content, list):
                    for item in content:
                        doc.add_paragraph(str(item), style='List Bullet')
                else:
                    doc.add_paragraph(str(content))

        doc.save(file_path)
        print(f"Model description saved to {file_path}")

    #TODO: export in ODD format

